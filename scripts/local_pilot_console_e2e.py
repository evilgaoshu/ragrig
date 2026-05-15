from __future__ import annotations

import argparse
import json
import os
import shutil
import socket
import subprocess
import tempfile
import threading
import time
from pathlib import Path
from typing import Any

import httpx
import uvicorn
from docx import Document

from ragrig.main import create_app
from scripts.local_pilot_smoke import _create_file_session_factory


class LocalPilotConsoleE2EError(RuntimeError):
    pass


def _find_free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def _write_minimal_text_pdf(path: Path, text: str) -> None:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 24 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]

    body = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body.extend(f"{index} 0 obj\n".encode("ascii"))
        body.extend(obj)
        body.extend(b"\nendobj\n")

    xref_offset = len(body)
    body.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    body.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        body.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    body.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(body))


def _create_fixture_files(root: Path) -> list[Path]:
    md_path = root / "pilot-console-e2e.md"
    md_path.write_text(
        "# Local Pilot Console E2E\n\n"
        "The Local Pilot E2E verifies upload, indexing, chunk preview, grounded answers, "
        "and citations from the Web Console.\n",
        encoding="utf-8",
    )

    pdf_path = root / "pilot-console-e2e.pdf"
    _write_minimal_text_pdf(
        pdf_path,
        "Local Pilot E2E PDF verifies parser coverage and indexing",
    )

    docx_path = root / "pilot-console-e2e.docx"
    document = Document()
    document.add_heading("Local Pilot E2E DOCX", level=1)
    document.add_paragraph("The DOCX fixture verifies browser upload and chunk indexing.")
    document.save(docx_path)

    return [md_path, pdf_path, docx_path]


def _wait_for_server(base_url: str, timeout_seconds: float) -> None:
    deadline = time.monotonic() + timeout_seconds
    last_error: Exception | None = None
    while time.monotonic() < deadline:
        try:
            response = httpx.get(f"{base_url}/health", timeout=1.0)
            if response.status_code == 200:
                return
        except Exception as exc:  # pragma: no cover - depends on process timing
            last_error = exc
        time.sleep(0.2)
    raise LocalPilotConsoleE2EError(f"server did not become ready: {last_error}")


def _run_playwright(
    *,
    base_url: str,
    files: list[Path],
    output: Path | None,
    headed: bool,
    timeout_ms: int,
) -> dict[str, Any]:
    if shutil.which("npx") is None:
        raise LocalPilotConsoleE2EError("npx is required to run the browser E2E check")
    if shutil.which("npm") is None:
        raise LocalPilotConsoleE2EError("npm is required to install the Playwright package")

    spec = Path(__file__).with_suffix(".mjs")
    package = os.environ.get("RAGRIG_PLAYWRIGHT_NPM_PACKAGE", "playwright@1.56.1")
    env = os.environ.copy()
    env.update(
        {
            "RAGRIG_CONSOLE_E2E_BASE_URL": base_url,
            "RAGRIG_CONSOLE_E2E_FILES": json.dumps([str(path) for path in files]),
            "RAGRIG_CONSOLE_E2E_HEADLESS": "0" if headed else "1",
            "RAGRIG_CONSOLE_E2E_TIMEOUT_MS": str(timeout_ms),
            "RAGRIG_CONSOLE_E2E_BROWSER_CHANNEL": os.environ.get(
                "RAGRIG_CONSOLE_E2E_BROWSER_CHANNEL", "chrome"
            ),
        }
    )
    if output is not None:
        output.parent.mkdir(parents=True, exist_ok=True)
        env["RAGRIG_CONSOLE_E2E_OUTPUT"] = str(output.resolve())

    with tempfile.TemporaryDirectory(prefix="ragrig-console-e2e-node-") as node_dir:
        node_root = Path(node_dir)
        node_spec = node_root / spec.name
        shutil.copy2(spec, node_spec)

        install_package = subprocess.run(
            ["npm", "install", "--no-save", "--no-package-lock", package],
            cwd=node_root,
            env=env,
            text=True,
            capture_output=True,
            check=False,
        )
        if install_package.returncode != 0:
            raise LocalPilotConsoleE2EError(
                "failed to install Playwright package:\n"
                f"{install_package.stdout}\n{install_package.stderr}"
            )

        command = ["node", str(node_spec)]
        result = subprocess.run(
            command,
            cwd=node_root,
            env=env,
            text=True,
            capture_output=True,
            check=False,
        )
        combined_output = f"{result.stdout}\n{result.stderr}"
        browser_missing = (
            "Executable doesn't exist" in combined_output
            or "Chromium distribution 'chrome' is not found" in combined_output
        )
        if result.returncode != 0 and browser_missing:
            install_browser = subprocess.run(
                ["npx", "playwright", "install", "chromium"],
                cwd=node_root,
                env=env,
                text=True,
                capture_output=True,
                check=False,
            )
            if install_browser.returncode != 0:
                raise LocalPilotConsoleE2EError(
                    "failed to install Playwright Chromium:\n"
                    f"{install_browser.stdout}\n{install_browser.stderr}"
                )
            retry_env = dict(env)
            retry_env["RAGRIG_CONSOLE_E2E_BROWSER_CHANNEL"] = ""
            result = subprocess.run(
                command,
                cwd=node_root,
                env=retry_env,
                text=True,
                capture_output=True,
                check=False,
            )

    if result.returncode != 0:
        raise LocalPilotConsoleE2EError(
            f"browser E2E failed:\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
        )

    if output is not None and output.exists():
        return json.loads(output.read_text(encoding="utf-8"))
    return json.loads(result.stdout)


def run_e2e(
    *,
    database_path: Path,
    output: Path | None = None,
    headed: bool = False,
    timeout_ms: int = 30000,
) -> dict[str, Any]:
    session_factory, engine = _create_file_session_factory(database_path)
    port = _find_free_port()
    base_url = f"http://127.0.0.1:{port}"
    app = create_app(check_database=lambda: None, session_factory=session_factory)
    server = uvicorn.Server(uvicorn.Config(app, host="127.0.0.1", port=port, log_level="warning"))
    thread = threading.Thread(target=server.run, daemon=True)
    thread.start()
    try:
        _wait_for_server(base_url, timeout_ms / 1000)
        with tempfile.TemporaryDirectory(prefix="ragrig-console-e2e-files-") as fixture_dir:
            files = _create_fixture_files(Path(fixture_dir))
            return _run_playwright(
                base_url=base_url,
                files=files,
                output=output,
                headed=headed,
                timeout_ms=timeout_ms,
            )
    finally:
        server.should_exit = True
        thread.join(timeout=10)
        engine.dispose()


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Run a browser-level Local Pilot Web Console E2E check."
    )
    parser.add_argument(
        "--database-path",
        type=Path,
        default=None,
        help="Optional SQLite database path. Defaults to a temporary file.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Optional JSON artifact output path.",
    )
    parser.add_argument(
        "--headed",
        action="store_true",
        help="Run Chromium headed for local debugging.",
    )
    parser.add_argument(
        "--timeout-ms",
        type=int,
        default=30000,
        help="Browser action timeout in milliseconds.",
    )
    return parser


def main() -> int:
    args = build_parser().parse_args()
    if args.database_path is not None:
        result = run_e2e(
            database_path=args.database_path,
            output=args.output,
            headed=args.headed,
            timeout_ms=args.timeout_ms,
        )
    else:
        with tempfile.TemporaryDirectory(prefix="ragrig-local-pilot-console-e2e-") as temp_dir:
            result = run_e2e(
                database_path=Path(temp_dir) / "local-pilot-console-e2e.db",
                output=args.output,
                headed=args.headed,
                timeout_ms=args.timeout_ms,
            )
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

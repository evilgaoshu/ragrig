from __future__ import annotations

from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.table import Table

from ragrig.parsers.base import ParseResult
from ragrig.parsers.sanitizer import sanitize_text_summary


class DocxParserError(ValueError):
    """Raised when a DOCX cannot be parsed by the Local Pilot parser."""


def _table_to_markdown(table: Table) -> str:
    """Convert a python-docx Table to a Markdown table string."""
    rows = table.rows
    if not rows:
        return ""
    cells: list[list[str]] = []
    for row in rows:
        cells.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
    if not cells:
        return ""
    header = cells[0]
    sep = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for row in cells[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


class DocxParser:
    parser_name = "docx"
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def parse(self, path: Path) -> ParseResult:
        raw_bytes = path.read_bytes()
        try:
            document = Document(path)
        except Exception as exc:
            raise DocxParserError(f"DOCX could not be read: {path.name}") from exc

        # Walk document body in order: paragraphs and tables interleaved
        parts: list[str] = []
        table_count = 0
        paragraph_count = 0
        for block in document.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
            if tag == "p":
                para_obj = next((p for p in document.paragraphs if p._element is block), None)
                if para_obj is not None:
                    text = para_obj.text.strip()
                    if text:
                        parts.append(text)
                        paragraph_count += 1
            elif tag == "tbl":
                tbl_obj = next((t for t in document.tables if t._element is block), None)
                if tbl_obj is not None:
                    md = _table_to_markdown(tbl_obj)
                    if md:
                        parts.append(md)
                        table_count += 1

        extracted_text = "\n".join(parts)
        if not extracted_text.strip():
            raise DocxParserError(f"DOCX has no extractable body text: {path.name}")

        summary, redactions = sanitize_text_summary(extracted_text)
        return ParseResult(
            extracted_text=extracted_text,
            content_hash=sha256(raw_bytes).hexdigest(),
            mime_type=self.mime_type,
            parser_name=self.parser_name,
            metadata={
                "parser_id": "parser.docx",
                "status": "success",
                "extension": path.suffix.lower(),
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "char_count": len(extracted_text),
                "byte_count": len(raw_bytes),
                "text_summary": summary,
                "redaction_count": redactions,
            },
        )

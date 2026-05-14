from __future__ import annotations

from hashlib import sha256
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from ragrig.ingestion.pipeline import _select_parser
from ragrig.parsers import DocxParser, PdfParser
from ragrig.parsers.docx import DocxParserError
from ragrig.parsers.pdf import PdfParserError

pytestmark = pytest.mark.unit


def _write_minimal_text_pdf(path: Path, text: str, *, include_blank_page: bool = False) -> None:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    page_kids = "[3 0 R 6 0 R]" if include_blank_page else "[3 0 R]"
    page_count = 2 if include_blank_page else 1
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids {page_kids} /Count {page_count} >>".encode("ascii"),
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream = f"BT /F1 24 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects.append(
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream"
    )
    if include_blank_page:
        objects.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>")

    body = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body.extend(f"{index} 0 obj\n".encode("ascii"))
        body.extend(obj)
        body.extend(b"\nendobj\n")

    xref_offset = len(body)
    body.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    body.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        body.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    body.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(body))


def test_pdf_parser_extracts_text_and_metadata(tmp_path) -> None:
    path = tmp_path / "guide.pdf"
    _write_minimal_text_pdf(path, "Local Pilot PDF")
    raw_bytes = path.read_bytes()

    result = PdfParser().parse(path)

    assert "Local Pilot PDF" in result.extracted_text
    assert result.content_hash == sha256(raw_bytes).hexdigest()
    assert result.mime_type == "application/pdf"
    assert result.parser_name == "pdf"
    assert result.metadata["parser_id"] == "parser.pdf"
    assert result.metadata["status"] == "success"
    assert result.metadata["extension"] == ".pdf"
    assert result.metadata["page_count"] == 1
    assert result.metadata["char_count"] == len(result.extracted_text)
    assert result.metadata["byte_count"] == len(raw_bytes)
    assert result.metadata["redaction_count"] == 0
    assert "Local Pilot PDF" in result.metadata["text_summary"]


def test_pdf_parser_raises_for_no_extractable_text(tmp_path) -> None:
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as file_obj:
        writer.write(file_obj)

    with pytest.raises(PdfParserError, match="no extractable text"):
        PdfParser().parse(path)


def test_pdf_parser_marks_pages_without_text_as_degraded(tmp_path) -> None:
    path = tmp_path / "partial.pdf"
    _write_minimal_text_pdf(path, "Extractable page", include_blank_page=True)

    result = PdfParser().parse(path)

    assert result.metadata["status"] == "degraded"
    assert result.metadata["page_count"] == 2
    assert result.metadata["degraded_pages"] == [2]
    assert "no extractable text" in result.metadata["degraded_reason"]


def test_pdf_parser_raises_for_encrypted_pdf(tmp_path) -> None:
    path = tmp_path / "encrypted.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.encrypt("secret")
    with path.open("wb") as file_obj:
        writer.write(file_obj)

    with pytest.raises(PdfParserError, match="encrypted"):
        PdfParser().parse(path)


def test_docx_parser_extracts_paragraphs_headings_lists_and_metadata(tmp_path) -> None:
    path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("Local Pilot Brief", level=1)
    document.add_paragraph("This is the opening paragraph.")
    document.add_paragraph("First list item", style="List Bullet")
    document.save(path)
    raw_bytes = path.read_bytes()

    result = DocxParser().parse(path)

    assert "Local Pilot Brief" in result.extracted_text
    assert "This is the opening paragraph." in result.extracted_text
    assert "First list item" in result.extracted_text
    assert result.content_hash == sha256(raw_bytes).hexdigest()
    assert (
        result.mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert result.parser_name == "docx"
    assert result.metadata["parser_id"] == "parser.docx"
    assert result.metadata["status"] == "success"
    assert result.metadata["extension"] == ".docx"
    assert result.metadata["paragraph_count"] == 3
    assert result.metadata["char_count"] == len(result.extracted_text)
    assert result.metadata["byte_count"] == len(raw_bytes)
    assert result.metadata["redaction_count"] == 0
    assert "Local Pilot Brief" in result.metadata["text_summary"]
    assert "limitations" in result.metadata


def test_docx_parser_raises_for_empty_body_text(tmp_path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    with pytest.raises(DocxParserError, match="no extractable body text"):
        DocxParser().parse(path)


def test_select_parser_uses_pdf_and_docx_parsers() -> None:
    assert _select_parser(Path("guide.pdf")).parser_name == "pdf"
    assert _select_parser(Path("brief.docx")).parser_name == "docx"
